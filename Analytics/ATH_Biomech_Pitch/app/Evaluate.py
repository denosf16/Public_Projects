import os
import pandas as pd
import numpy as np
import zipfile
import io
import joblib
import matplotlib.pyplot as plt
import seaborn as sns
import scipy.stats as stats
import streamlit as st
from docx import Document
from docx.shared import Inches
from sklearn.metrics import (
    mean_squared_error,
    mean_absolute_error,
    median_absolute_error,
    r2_score
)
from utils import config

# ------------------------------
# 🔹 Paths
# ------------------------------
plot_base = os.path.join(config.OUTPUT_PATH, "plots", "diagnostics")
metrics_path = os.path.join(config.OUTPUT_PATH, "models", "model_metrics.csv")

# ------------------------------
# 🔷 Delta Gradient Formatter
# ------------------------------
def highlight_delta(val):
    if isinstance(val, str) or pd.isna(val):
        return ""
    max_val = 0.3
    norm_val = min(abs(val), max_val) / max_val
    if val > 0:
        green = int(255 - norm_val * 120)
        return f"background-color: rgb(200, 255, {green}); color: #000000;"
    elif val < 0:
        red = int(255 - norm_val * 120)
        return f"background-color: rgb(255, {red}, {red}); color: #000000;"
    else:
        return "background-color: #ffffff; color: #000000;"

# ------------------------------
# 📄 Generate Word Report
# ------------------------------
def generate_word_report(row, img_dir, model_name):
    doc = Document()
    doc.add_heading(f'Model Evaluation Report: {model_name}', 0)

    doc.add_paragraph("This report summarizes the performance of the selected model evaluated on a held-out test set using:")
    doc.add_paragraph("- Root Mean Squared Error (RMSE)")
    doc.add_paragraph("- Mean Absolute Error (MAE)")
    doc.add_paragraph("- R² (Coefficient of Determination)")

    doc.add_heading("Metrics", level=1)
    doc.add_paragraph(f"Validation R²: {row['Val_R2']:.2f}")
    doc.add_paragraph(f"Validation RMSE: {row['Val_RMSE']:.2f}")
    doc.add_paragraph(f"Validation MAE: {row['Val_MAE']:.2f}")
    doc.add_paragraph(f"Test R²: {row['Test_R2']:.2f}")
    doc.add_paragraph(f"Test RMSE: {row['Test_RMSE']:.2f}")
    doc.add_paragraph(f"Test MAE: {row['Test_MAE']:.2f}")

    doc.add_heading("Diagnostic Plots", level=1)
    for img in ["pred_vs_actual.png", "residuals.png", "qq_plot.png", "residuals_vs_fitted.png"]:
        img_path = os.path.join(img_dir, img)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
            doc.add_paragraph(img.replace(".png", "").replace("_", " ").title())
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------------------
# 🔷 Streamlit Page
# ------------------------------
def show():
    st.title("📊 Model Evaluation")

    with st.expander("📘 Evaluation Summary", expanded=True):
        st.markdown("""
        Model performance is assessed on a **held-out test set** using three standard regression metrics:

        - **Root Mean Squared Error (RMSE)**: Penalizes larger errors more heavily. Lower is better.  
        - **Mean Absolute Error (MAE)**: Easier to interpret average error. Lower is better.  
        - **R² (Coefficient of Determination)**: Proportion of variance explained. Higher is better.

        Each model is visualized with:
        - **Predicted vs. Actual Plot**: Alignment of predicted and true pitch speeds.  
        - **Residuals vs. Predicted Plot**: Detects trends or heteroscedasticity in errors.

        **Key Insight**:  
        The **Full Model** achieved the best performance, explaining over **73% of pitch velocity variance**. This reinforces the value of integrating upper-body, lower-body, and ground-force features together.
        """)

    try:
        metrics_df = pd.read_csv(metrics_path)
        numeric_cols = [col for col in metrics_df.columns if "R2" in col or "RMSE" in col or "MAE" in col]
        metrics_df[numeric_cols] = metrics_df[numeric_cols].round(2)

        display_cols = ["Technique_Model", "Val_R2", "Val_RMSE", "Val_MAE", "Test_R2", "Test_RMSE", "Test_MAE"]
        metrics_df_sorted = metrics_df.sort_values("Test_R2", ascending=False)

        st.markdown("### 🧾 Evaluation Summary")
        st.dataframe(metrics_df_sorted[display_cols], use_container_width=True)

        csv_export = metrics_df_sorted[display_cols].to_csv(index=False)
        st.download_button(
            label="⬇️ Download Metrics CSV",
            data=csv_export,
            file_name="model_evaluation_summary.csv",
            mime="text/csv"
        )

        st.markdown("---")
        st.markdown("### 🔍 Side-by-Side Model Comparison")

        options = metrics_df_sorted["Technique_Model"].tolist()
        model_1 = st.selectbox("Select First Model", options, index=0)
        model_2 = st.selectbox("Select Second Model", options, index=1)

        df1 = metrics_df[metrics_df["Technique_Model"] == model_1].iloc[0]
        df2 = metrics_df[metrics_df["Technique_Model"] == model_2].iloc[0]

        delta_df = pd.DataFrame({
            "Metric": ["Test_R2", "Test_RMSE", "Test_MAE", "Val_R2"],
            model_1: [df1["Test_R2"], df1["Test_RMSE"], df1["Test_MAE"], df1["Val_R2"]],
            model_2: [df2["Test_R2"], df2["Test_RMSE"], df2["Test_MAE"], df2["Val_R2"]],
            "Delta": [
                round(df1["Test_R2"] - df2["Test_R2"], 2),
                round(df1["Test_RMSE"] - df2["Test_RMSE"], 2),
                round(df1["Test_MAE"] - df2["Test_MAE"], 2),
                round(df1["Val_R2"] - df2["Val_R2"], 2)
            ]
        })

        st.dataframe(delta_df.style.map(highlight_delta, subset=["Delta"]), use_container_width=True)

        st.markdown("### 🖼️ Diagnostic Plots")

        col1, col2 = st.columns(2)
        for i, model_id in enumerate([model_1, model_2]):
            technique, variant = model_id.split("_", 1)
            img_dir = os.path.join(plot_base, variant, technique)

            if os.path.exists(img_dir):
                with (col1 if i == 0 else col2):
                    st.markdown(f"#### {model_id}")
                    for img_file in ["pred_vs_actual.png", "residuals.png", "qq_plot.png", "residuals_vs_fitted.png"]:
                        img_path = os.path.join(img_dir, img_file)
                        if os.path.exists(img_path):
                            st.image(img_path, caption=img_file.replace(".png", "").replace("_", " ").title(), use_container_width=True)

                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w") as zipf:
                        for img_file in ["pred_vs_actual.png", "residuals.png", "qq_plot.png", "residuals_vs_fitted.png"]:
                            full_path = os.path.join(img_dir, img_file)
                            if os.path.exists(full_path):
                                zipf.write(full_path, arcname=img_file)
                    zip_buffer.seek(0)

                    st.download_button(
                        label=f"⬇️ Download Plots for {model_id}",
                        data=zip_buffer,
                        file_name=f"{model_id}_diagnostics.zip",
                        mime="application/zip",
                        key=f"zip_{model_id}"
                    )

                    if st.button(f"📄 Generate Report for {model_id}", key=f"report_{model_id}"):
                        buffer = generate_word_report(df1 if i == 0 else df2, img_dir, model_id)
                        st.download_button(
                            label="⬇️ Download Word Report",
                            data=buffer,
                            file_name=f"{model_id}_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{model_id}"
                        )

    except Exception as e:
        st.error(f"❌ Could not load evaluation data: {e}")






